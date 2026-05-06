"""
PDF Intelligence Converter - Premium Batsirai Design
Production Version 4.1.0 - Enhanced with Batch Processing, OCR Language Detection, Mobile Responsiveness, Notifications & DB Migrations
Run with: streamlit run app.py
"""

import streamlit as st
import pandas as pd
import tempfile
import os
from pathlib import Path
import pdfplumber
import fitz  # PyMuPDF
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
from PIL import Image
import io
from datetime import datetime, timezone
import json
import csv
import sqlite3
import hashlib
from typing import List, Dict, Any
import zipfile
import re
import time
import warnings
import logging
import gc
from collections import deque
from functools import wraps

warnings.filterwarnings("ignore")
logging.getLogger('streamlit').setLevel(logging.ERROR)

OCR_AVAILABLE = False
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    from langdetect import detect, DetectorFactory
    DetectorFactory.seed = 0
    OCR_AVAILABLE = True
except ImportError:
    pass

APP_VERSION = "4.1.0"
APP_NAME = "PDF Intelligence Converter"
DEPLOYMENT_MODE = os.environ.get("DEPLOYMENT_MODE", "production")
SESSION_TIMEOUT_MINUTES = 60
MAX_BATCH_SIZE = 20
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

# ============================================
# NOTIFICATION SYSTEM
# ============================================
class NotificationSystem:
    """Enhanced notification system with persistence and categories"""
    
    def __init__(self):
        if "notifications" not in st.session_state:
            st.session_state.notifications = deque(maxlen=50)  # Keep last 50 notifications
        if "unread_count" not in st.session_state:
            st.session_state.unread_count = 0
    
    def add(self, message: str, type: str = "info", duration: int = 5, action_url: str = None):
        """Add a new notification
        
        Args:
            message: Notification message
            type: info, success, warning, error
            duration: Auto-dismiss seconds (0 for persistent)
            action_url: Optional action URL
        """
        notification = {
            "id": hashlib.md5(f"{message}{datetime.now()}".encode()).hexdigest()[:8],
            "message": message,
            "type": type,
            "timestamp": datetime.now(),
            "read": False,
            "duration": duration,
            "action_url": action_url
        }
        st.session_state.notifications.appendleft(notification)
        st.session_state.unread_count += 1
        
        # Also show in Streamlit for immediate feedback
        if type == "success":
            st.success(message)
        elif type == "error":
            st.error(message)
        elif type == "warning":
            st.warning(message)
        else:
            st.info(message)
        
        return notification["id"]
    
    def mark_as_read(self, notification_id: str):
        """Mark a notification as read"""
        for notif in st.session_state.notifications:
            if notif["id"] == notification_id:
                if not notif["read"]:
                    notif["read"] = True
                    st.session_state.unread_count = max(0, st.session_state.unread_count - 1)
                break
    
    def mark_all_read(self):
        """Mark all notifications as read"""
        for notif in st.session_state.notifications:
            notif["read"] = True
        st.session_state.unread_count = 0
    
    def clear_all(self):
        """Clear all notifications"""
        st.session_state.notifications.clear()
        st.session_state.unread_count = 0
    
    def get_unread_count(self) -> int:
        """Get number of unread notifications"""
        return st.session_state.unread_count
    
    def render_notification_center(self):
        """Render notification center UI"""
        if not st.session_state.notifications:
            st.info("📭 No notifications")
            return
        
        # Header with actions
        col1, col2, col3 = st.columns([2, 1, 1])
        with col1:
            st.markdown(f"**📬 Notifications** ({len([n for n in st.session_state.notifications if not n['read']])} unread)")
        with col2:
            if st.button("✓ Mark all read", key="mark_all_read", use_container_width=True):
                self.mark_all_read()
                st.rerun()
        with col3:
            if st.button("🗑 Clear all", key="clear_all_notif", use_container_width=True):
                self.clear_all()
                st.rerun()
        
        st.markdown("---")
        
        # Display notifications
        for notif in st.session_state.notifications:
            # Style based on type and read status
            opacity = "0.7" if notif["read"] else "1"
            bg_color = {
                "success": "#d4edda" if st.session_state.theme == "light" else "#155724",
                "error": "#f8d7da" if st.session_state.theme == "light" else "#721c24",
                "warning": "#fff3cd" if st.session_state.theme == "light" else "#856404",
                "info": "#d1ecf1" if st.session_state.theme == "light" else "#0c5460"
            }.get(notif["type"], "#e2e3e5" if st.session_state.theme == "light" else "#383d41")
            
            text_color = "#000000" if st.session_state.theme == "light" else "#ffffff"
            
            with st.container():
                st.markdown(f"""
                <div style="
                    background: {bg_color};
                    padding: 12px;
                    border-radius: 8px;
                    margin-bottom: 8px;
                    opacity: {opacity};
                    border-left: 4px solid {'#' + hashlib.md5(notif['type'].encode()).hexdigest()[:6]};
                ">
                    <div style="display: flex; justify-content: space-between; align-items: center;">
                        <div style="flex: 1;">
                            <strong>{notif['type'].upper()}</strong><br>
                            {notif['message']}
                            <div style="font-size: 11px; color: {text_color}99; margin-top: 4px;">
                                {notif['timestamp'].strftime('%Y-%m-%d %H:%M:%S')}
                            </div>
                        </div>
                        <div style="display: flex; gap: 8px;">
                            {f'<a href="{notif["action_url"]}" target="_blank" style="text-decoration: none;">🔗</a>' if notif["action_url"] else ''}
                            {f'<button onclick="mark_read_{notif["id"]}">✓</button>' if not notif["read"] else '✓ Read'}
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                if not notif["read"]:
                    if st.button(f"Mark read", key=f"read_{notif['id']}", use_container_width=True):
                        self.mark_as_read(notif["id"])
                        st.rerun()
            
            st.markdown("<br>", unsafe_allow_html=True)

# ============================================
# DATABASE MIGRATIONS
# ============================================
class DatabaseManager:
    """Handle database schema migrations and versioning"""
    
    def __init__(self, db_path='users.db'):
        self.db_path = db_path
        self.current_version = 2  # Increment when schema changes
        
    def get_db_version(self):
        """Get current database schema version"""
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        
        try:
            c.execute("SELECT version FROM db_version ORDER BY version DESC LIMIT 1")
            version = c.fetchone()
            if version:
                return version[0]
        except sqlite3.OperationalError:
            # Version table doesn't exist yet
            pass
        finally:
            conn.close()
        return 0
    
    def set_db_version(self, version):
        """Set database schema version"""
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        c.execute("CREATE TABLE IF NOT EXISTS db_version (version INTEGER, applied_at TIMESTAMP)")
        c.execute("INSERT INTO db_version (version, applied_at) VALUES (?, ?)", 
                 (version, datetime.now()))
        conn.commit()
        conn.close()
    
    def migrate_v1_to_v2(self):
        """Migrate from v1 to v2 - Add new columns and indexes"""
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        
        try:
            # Add new columns to conversion_history if they don't exist
            c.execute("PRAGMA table_info(conversion_history)")
            columns = [col[1] for col in c.fetchall()]
            
            new_columns = {
                "processing_time": "REAL",
                "error_message": "TEXT",
                "file_pages": "INTEGER",
                "batch_id": "TEXT",
                "conversion_status": "TEXT DEFAULT 'completed'"
            }
            
            for col_name, col_type in new_columns.items():
                if col_name not in columns:
                    c.execute(f"ALTER TABLE conversion_history ADD COLUMN {col_name} {col_type}")
            
            # Add indexes for better query performance
            c.execute("CREATE INDEX IF NOT EXISTS idx_timestamp ON conversion_history(timestamp)")
            c.execute("CREATE INDEX IF NOT EXISTS idx_username_timestamp ON conversion_history(username, timestamp)")
            c.execute("CREATE INDEX IF NOT EXISTS idx_batch_id ON conversion_history(batch_id)")
            c.execute("CREATE INDEX IF NOT EXISTS idx_conversion_status ON conversion_history(conversion_status)")
            
            # Add new table for batch jobs
            c.execute("""
                CREATE TABLE IF NOT EXISTS batch_jobs (
                    batch_id TEXT PRIMARY KEY,
                    username TEXT,
                    created_at TIMESTAMP,
                    total_files INTEGER,
                    completed_files INTEGER DEFAULT 0,
                    failed_files INTEGER DEFAULT 0,
                    status TEXT DEFAULT 'processing',
                    FOREIGN KEY (username) REFERENCES users(username)
                )
            """)
            
            # Add user preferences table
            c.execute("""
                CREATE TABLE IF NOT EXISTS user_preferences (
                    username TEXT PRIMARY KEY,
                    default_output_format TEXT DEFAULT 'Excel (XLSX)',
                    default_extraction_mode TEXT DEFAULT 'Smart (Text + Tables)',
                    auto_convert BOOLEAN DEFAULT 0,
                    notify_on_completion BOOLEAN DEFAULT 1,
                    theme TEXT DEFAULT 'light',
                    FOREIGN KEY (username) REFERENCES users(username)
                )
            """)
            
            # Add user activity tracking
            c.execute("""
                CREATE TABLE IF NOT EXISTS user_activity (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    username TEXT,
                    action TEXT,
                    details TEXT,
                    timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    ip_address TEXT
                )
            """)
            
            conn.commit()
            
            # Update version
            self.set_db_version(2)
            return True
            
        except Exception as e:
            print(f"Migration error: {e}")
            conn.rollback()
            return False
        finally:
            conn.close()
    
    def run_migrations(self):
        """Run all pending migrations"""
        current_version = self.get_db_version()
        
        if current_version < 1:
            self.init_database()
        
        if current_version < 2:
            success = self.migrate_v1_to_v2()
            if success:
                print("Successfully migrated to version 2")
    
    def init_database(self):
        """Initialize database with initial schema"""
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        
        # Users table
        c.execute('''CREATE TABLE IF NOT EXISTS users
                     (username TEXT PRIMARY KEY, 
                      password_hash TEXT, 
                      role TEXT DEFAULT 'user',
                      created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                      last_login TIMESTAMP)''')
        
        # Conversion history table
        c.execute('''CREATE TABLE IF NOT EXISTS conversion_history
                     (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                      username TEXT,
                      filename TEXT,
                      output_format TEXT,
                      timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                      file_size INTEGER)''')
        
        # Create admin user
        c.execute("SELECT * FROM users WHERE username = 'admin'")
        if not c.fetchone():
            admin_hash = hashlib.sha256("admin123".encode()).hexdigest()
            c.execute("INSERT INTO users (username, password_hash, role, created_at) VALUES (?, ?, ?, ?)",
                     ('admin', admin_hash, 'admin', datetime.now()))
        
        conn.commit()
        conn.close()
        
        # Set version
        self.set_db_version(1)

# Initialize database with migrations
db_manager = DatabaseManager()
db_manager.run_migrations()

# Initialize notification system
notification_system = NotificationSystem()

# ============================================
# BATCH PROCESSING SYSTEM
# ============================================
class BatchProcessor:
    """Handle batch processing of multiple PDF files"""
    
    def __init__(self):
        if "batch_jobs" not in st.session_state:
            st.session_state.batch_jobs = {}
        if "active_batch" not in st.session_state:
            st.session_state.active_batch = None
    
    def create_batch_job(self, files: List, output_format: str, extraction_mode: str, 
                        extraction_options: Dict) -> str:
        """Create a new batch processing job"""
        batch_id = hashlib.md5(f"{st.session_state.username}{datetime.now()}".encode()).hexdigest()[:12]
        
        job = {
            "batch_id": batch_id,
            "username": st.session_state.username,
            "files": [{"name": f.name, "bytes": f.getvalue(), "size": len(f.getvalue())} for f in files],
            "total_files": len(files),
            "completed_files": 0,
            "failed_files": 0,
            "status": "pending",
            "output_format": output_format,
            "extraction_mode": extraction_mode,
            "extraction_options": extraction_options,
            "created_at": datetime.now(),
            "results": [],
            "progress": 0
        }
        
        st.session_state.batch_jobs[batch_id] = job
        st.session_state.active_batch = batch_id
        
        # Save to database
        self._save_batch_to_db(job)
        
        notification_system.add(
            f"Batch job created with {len(files)} files. Processing started.",
            "info",
            duration=3
        )
        
        return batch_id
    
    def _save_batch_to_db(self, job):
        """Save batch job to database"""
        conn = sqlite3.connect('users.db')
        c = conn.cursor()
        c.execute("""
            INSERT OR REPLACE INTO batch_jobs 
            (batch_id, username, created_at, total_files, completed_files, failed_files, status)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (job["batch_id"], job["username"], job["created_at"], 
              job["total_files"], job["completed_files"], job["failed_files"], job["status"]))
        conn.commit()
        conn.close()
    
    def process_batch(self, batch_id: str, progress_callback=None):
        """Process all files in a batch job"""
        if batch_id not in st.session_state.batch_jobs:
            return False
        
        job = st.session_state.batch_jobs[batch_id]
        job["status"] = "processing"
        
        results = []
        
        for idx, file_data in enumerate(job["files"]):
            try:
                start_time = time.time()
                
                # Process the file
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                    tmp_file.write(file_data["bytes"])
                    tmp_path = tmp_file.name
                
                # Extract content
                content = extract_pdf_intelligent(
                    tmp_path, 
                    job["extraction_mode"], 
                    job["extraction_options"].get("extract_tables", True)
                )
                
                # Convert based on output format
                output_buffer = self._convert_content(content, job["output_format"], job["extraction_options"])
                
                processing_time = time.time() - start_time
                
                result = {
                    "filename": file_data["name"],
                    "success": True,
                    "output": output_buffer,
                    "processing_time": processing_time,
                    "pages": content["pages"],
                    "tables": len(content["tables"])
                }
                
                # Save to history
                save_conversion_history(
                    st.session_state.username, 
                    file_data["name"], 
                    job["output_format"], 
                    file_data["size"],
                    processing_time,
                    content["pages"],
                    batch_id
                )
                
                job["completed_files"] += 1
                results.append(result)
                
                os.unlink(tmp_path)
                
            except Exception as e:
                job["failed_files"] += 1
                results.append({
                    "filename": file_data["name"],
                    "success": False,
                    "error": str(e)
                })
                
                # Log error to database
                save_conversion_history(
                    st.session_state.username, 
                    file_data["name"], 
                    job["output_format"], 
                    file_data["size"],
                    None,
                    None,
                    batch_id,
                    status="failed",
                    error=str(e)
                )
            
            # Update progress
            job["progress"] = ((job["completed_files"] + job["failed_files"]) / job["total_files"]) * 100
            if progress_callback:
                progress_callback(job["progress"])
            
            self._save_batch_to_db(job)
        
        job["results"] = results
        job["status"] = "completed" if job["failed_files"] == 0 else "partial"
        
        # Final notification
        if job["failed_files"] == 0:
            notification_system.add(
                f"Batch processing complete! {job['completed_files']} files converted successfully.",
                "success",
                duration=5,
                action_url="#history"
            )
        else:
            notification_system.add(
                f"Batch completed with {job['failed_files']} failures out of {job['total_files']} files.",
                "warning",
                duration=5
            )
        
        return True
    
    def _convert_content(self, content, output_format, options):
        """Convert extracted content to desired format"""
        result_buffer = io.BytesIO()
        
        if output_format == "Excel (XLSX)":
            convert_to_excel(content, result_buffer, options.get("include_metadata", True))
        elif output_format == "Word (DOCX)":
            convert_to_word(content, result_buffer)
        elif output_format == "Text (TXT)":
            text_content = "\n\n".join([p["content"] for p in content["text"]])
            result_buffer = io.BytesIO(text_content.encode('utf-8'))
        elif output_format == "CSV":
            text_stream = io.TextIOWrapper(result_buffer, 'utf-8', newline='')
            writer = csv.writer(text_stream)
            writer.writerow(["Page", "Content"])
            for page in content["text"]:
                writer.writerow([page["page"], page["content"]])
            text_stream.flush()
        elif output_format == "JSON":
            json_content = {
                "metadata": content["metadata"], 
                "pages": content["pages"], 
                "text": [{"page": p["page"], "content": p["content"]} for p in content["text"]], 
                "tables": [{"page": t["page"], "table": t["table"]} for t in content["tables"]]
            }
            result_buffer = io.BytesIO(json.dumps(json_content, indent=2, ensure_ascii=False).encode('utf-8'))
        elif output_format == "Markdown":
            md = convert_to_markdown(content)
            result_buffer = io.BytesIO(md.encode('utf-8'))
        elif output_format == "HTML":
            html_content = convert_to_html(content)
            result_buffer = io.BytesIO(html_content.encode('utf-8'))
        
        result_buffer.seek(0)
        return result_buffer
    
    def get_batch_results_zip(self, batch_id: str) -> io.BytesIO:
        """Create a ZIP file with all batch conversion results"""
        if batch_id not in st.session_state.batch_jobs:
            return None
        
        job = st.session_state.batch_jobs[batch_id]
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for result in job["results"]:
                if result["success"]:
                    ext_map = {
                        "Excel (XLSX)": ".xlsx",
                        "Word (DOCX)": ".docx",
                        "Text (TXT)": ".txt",
                        "CSV": ".csv",
                        "JSON": ".json",
                        "Markdown": ".md",
                        "HTML": ".html"
                    }
                    ext = ext_map.get(job["output_format"], ".txt")
                    filename = Path(result["filename"]).stem + ext
                    zip_file.writestr(filename, result["output"].getvalue())
        
        zip_buffer.seek(0)
        return zip_buffer
    
    def render_batch_status(self, batch_id: str):
        """Render batch processing status UI"""
        if batch_id not in st.session_state.batch_jobs:
            return
        
        job = st.session_state.batch_jobs[batch_id]
        
        with st.container():
            st.markdown(f"### 📦 Batch Job: {batch_id}")
            
            # Progress bar
            progress = job["progress"]
            st.progress(progress / 100)
            
            # Stats
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Total Files", job["total_files"])
            with col2:
                st.metric("Completed", job["completed_files"])
            with col3:
                st.metric("Failed", job["failed_files"])
            with col4:
                st.metric("Status", job["status"].upper())
            
            # Results table if completed
            if job["status"] in ["completed", "partial"]:
                results_df = pd.DataFrame([
                    {
                        "File": r["filename"],
                        "Status": "✅ Success" if r["success"] else "❌ Failed",
                        "Pages": r.get("pages", "N/A"),
                        "Tables": r.get("tables", "N/A"),
                        "Time": f"{r.get('processing_time', 0):.2f}s" if r.get("processing_time") else "N/A"
                    }
                    for r in job["results"]
                ])
                st.dataframe(results_df, use_container_width=True)
                
                # Download button
                if job["completed_files"] > 0:
                    zip_buffer = self.get_batch_results_zip(batch_id)
                    if zip_buffer:
                        st.download_button(
                            label=f"📥 Download All Results ({job['completed_files']} files)",
                            data=zip_buffer.getvalue(),
                            file_name=f"batch_{batch_id}_results.zip",
                            mime="application/zip",
                            use_container_width=True
                        )

# ============================================
# OCR LANGUAGE DETECTION
# ============================================
class OCRProcessor:
    """Enhanced OCR with automatic language detection"""
    
    def __init__(self):
        self.supported_languages = {
            'eng': 'English',
            'spa': 'Spanish',
            'fra': 'French',
            'deu': 'German',
            'ita': 'Italian',
            'por': 'Portuguese',
            'rus': 'Russian',
            'jpn': 'Japanese',
            'kor': 'Korean',
            'chi_sim': 'Chinese (Simplified)',
            'chi_tra': 'Chinese (Traditional)',
            'ara': 'Arabic',
            'hin': 'Hindi',
            'nld': 'Dutch',
            'pol': 'Polish',
            'tur': 'Turkish',
            'swe': 'Swedish',
            'fin': 'Finnish',
            'nor': 'Norwegian',
            'dan': 'Danish'
        }
        
        if not OCR_AVAILABLE:
            return
    
    def detect_language_from_pdf(self, pdf_bytes: bytes, sample_pages: int = 3) -> str:
        """Detect the primary language of a PDF document"""
        if not OCR_AVAILABLE:
            return 'eng'
        
        try:
            # Convert first few pages to images
            images = convert_from_bytes(pdf_bytes, first_page=1, last_page=sample_pages, dpi=100)
            
            text_samples = []
            for img in images:
                # Perform OCR with default English to get text sample
                text = pytesseract.image_to_string(img, lang='eng')
                if len(text.strip()) > 50:  # Only use if we got substantial text
                    text_samples.append(text)
            
            if not text_samples:
                return 'eng'
            
            # Detect language from combined text
            combined_text = " ".join(text_samples)
            detected_lang = detect(combined_text)
            
            # Map detected language to Tesseract language code
            lang_map = {
                'en': 'eng',
                'es': 'spa',
                'fr': 'fra',
                'de': 'deu',
                'it': 'ita',
                'pt': 'por',
                'ru': 'rus',
                'ja': 'jpn',
                'ko': 'kor',
                'zh-cn': 'chi_sim',
                'zh-tw': 'chi_tra',
                'ar': 'ara',
                'hi': 'hin',
                'nl': 'nld',
                'pl': 'pol',
                'tr': 'tur',
                'sv': 'swe',
                'fi': 'fin',
                'no': 'nor',
                'da': 'dan'
            }
            
            return lang_map.get(detected_lang[:2], 'eng')
            
        except Exception as e:
            notification_system.add(f"Language detection failed: {str(e)}", "warning", duration=2)
            return 'eng'
    
    def perform_ocr(self, pdf_bytes: bytes, languages: List[str] = None, dpi: int = 200) -> Dict:
        """Perform OCR with language detection"""
        if not OCR_AVAILABLE:
            return {"text": "", "error": "OCR not available"}
        
        if languages is None:
            languages = ['eng']
        
        try:
            # Convert PDF to images
            images = convert_from_bytes(pdf_bytes, dpi=dpi)
            
            all_text = []
            total_pages = len(images)
            
            # Progress tracking
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for page_num, image in enumerate(images, 1):
                status_text.text(f"Processing page {page_num}/{total_pages}...")
                
                # Try with specified languages
                lang_str = '+'.join(languages)
                page_text = pytesseract.image_to_string(image, lang=lang_str)
                all_text.append({
                    'page': page_num,
                    'text': page_text,
                    'length': len(page_text)
                })
                
                progress_bar.progress(page_num / total_pages)
            
            status_text.empty()
            progress_bar.empty()
            
            combined_text = "\n\n".join([p['text'] for p in all_text])
            
            return {
                "text": combined_text,
                "pages": all_text,
                "total_pages": total_pages,
                "total_chars": len(combined_text),
                "total_words": len(combined_text.split()),
                "languages_used": languages
            }
            
        except Exception as e:
            return {"text": "", "error": str(e)}
    
    def render_ocr_interface(self):
        """Render OCR interface with language detection"""
        st.markdown("### 🔍 OCR with Language Detection")
        
        uploaded_file = st.file_uploader("Upload PDF for OCR", type=['pdf'], key="ocr_upload")
        
        if uploaded_file:
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("🔮 Auto-detect Language", use_container_width=True):
                    with st.spinner("Detecting language..."):
                        detected_lang = self.detect_language_from_pdf(uploaded_file.getvalue())
                        st.session_state.detected_lang = detected_lang
                        notification_system.add(
                            f"Detected language: {self.supported_languages.get(detected_lang, detected_lang)}",
                            "success",
                            duration=3
                        )
            
            with col2:
                # Manual language selection
                available_langs = list(self.supported_languages.keys())
                selected_langs = st.multiselect(
                    "Or select languages manually",
                    options=available_langs,
                    format_func=lambda x: self.supported_languages[x],
                    default=[st.session_state.get('detected_lang', 'eng')]
                )
            
            dpi = st.slider("OCR Quality (DPI)", 100, 300, 200, help="Higher DPI = better accuracy but slower")
            
            if st.button("🚀 Start OCR", use_container_width=True):
                with st.spinner("Performing OCR..."):
                    result = self.perform_ocr(uploaded_file.getvalue(), selected_langs or ['eng'], dpi)
                    
                    if "error" in result:
                        st.error(f"OCR failed: {result['error']}")
                    else:
                        st.success(f"OCR complete! Extracted {result['total_chars']} characters from {result['total_pages']} pages")
                        
                        # Display statistics
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Pages", result['total_pages'])
                        with col2:
                            st.metric("Characters", result['total_chars'])
                        with col3:
                            st.metric("Words", result['total_words'])
                        
                        # Show extracted text
                        with st.expander("View Extracted Text"):
                            st.text_area("OCR Result", result['text'], height=400)
                        
                        # Download results
                        download_buffer = io.BytesIO(result['text'].encode('utf-8'))
                        st.download_button(
                            label="📥 Download OCR Text",
                            data=download_buffer.getvalue(),
                            file_name=f"ocr_result_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain",
                            use_container_width=True
                        )
                        
                        notification_system.add(
                            f"OCR completed successfully for {uploaded_file.name}",
                            "success",
                            duration=3
                        )

# ============================================
# MOBILE RESPONSIVENESS
# ============================================
class MobileOptimizer:
    """Handle mobile device detection and responsive UI"""
    
    def __init__(self):
        self.is_mobile = self._detect_mobile()
        if self.is_mobile:
            self._apply_mobile_styles()
    
    def _detect_mobile(self) -> bool:
        """Detect if user is on mobile device"""
        try:
            # Check viewport width via JavaScript (if available)
            # Fallback to user agent detection
            user_agent = st.get_option("browser.userAgent", "")
            mobile_keywords = ['Android', 'webOS', 'iPhone', 'iPad', 'iPod', 'BlackBerry', 'Windows Phone', 'Mobile']
            return any(keyword in user_agent for keyword in mobile_keywords)
        except:
            return False
    
    def _apply_mobile_styles(self):
        """Apply mobile-specific CSS styles"""
        st.markdown("""
        <style>
        /* Mobile responsiveness */
        @media (max-width: 768px) {
            .block-container {
                padding: 0.5rem !important;
            }
            
            .card, .card-soft, .hero {
                padding: 16px !important;
                margin-bottom: 10px !important;
            }
            
            .title, .login-title {
                font-size: 20px !important;
            }
            
            .metric-v {
                font-size: 20px !important;
            }
            
            .metric-k {
                font-size: 10px !important;
            }
            
            .chip-container {
                gap: 6px !important;
            }
            
            .chip {
                padding: 4px 10px !important;
                font-size: 10px !important;
            }
            
            .stTabs [data-baseweb="tab"] {
                padding: 8px 12px !important;
                font-size: 12px !important;
            }
            
            div.stButton > button {
                padding: 8px 16px !important;
                font-size: 13px !important;
                min-height: 36px !important;
            }
            
            [data-testid="stDataFrame"] {
                font-size: 12px !important;
                overflow-x: auto !important;
            }
            
            /* Make tables scrollable on mobile */
            .stDataFrame {
                max-width: 100vw;
                overflow-x: auto;
            }
            
            /* Adjust column layout for mobile */
            .row-widget.stColumns {
                flex-direction: column !important;
            }
            
            /* Larger touch targets */
            button, [role="button"], .stDownloadButton button {
                min-height: 44px !important;
            }
        }
        
        /* Touch-friendly spacing */
        @media (hover: none) and (pointer: coarse) {
            .stButton button, 
            .stDownloadButton button,
            .stSelectbox div[role="button"] {
                min-height: 44px !important;
                padding: 12px !important;
            }
            
            input, select, textarea {
                font-size: 16px !important; /* Prevent zoom on focus */
            }
        }
        </style>
        """, unsafe_allow_html=True)
    
    def render_mobile_notice(self):
        """Render notice for mobile users"""
        if self.is_mobile:
            with st.sidebar:
                st.info("📱 Mobile Mode Active • Optimized for touch")
    
    def get_layout_columns(self, desktop_cols: tuple, mobile_cols: tuple = (1,)) -> tuple:
        """Return appropriate column layout based on device"""
        if self.is_mobile:
            return mobile_cols
        return desktop_cols
    
    def get_grid_columns(self, items: List, desktop_cols: int = 3, mobile_cols: int = 1) -> List[List]:
        """Split items into grid columns for responsive layout"""
        cols_per_row = 1 if self.is_mobile else desktop_cols
        
        grid = []
        for i in range(0, len(items), cols_per_row):
            grid.append(items[i:i + cols_per_row])
        return grid

# ============================================
# ENHANCED DATABASE FUNCTIONS
# ============================================
def save_conversion_history(username, filename, output_format, file_size, processing_time=None, 
                           file_pages=None, batch_id=None, status="completed", error=None):
    """Enhanced save function with additional fields"""
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    
    try:
        # Check if column exists
        c.execute("SELECT processing_time FROM conversion_history LIMIT 1")
    except sqlite3.OperationalError:
        # Column doesn't exist - run migration
        db_manager.migrate_v1_to_v2()
    
    c.execute("""
        INSERT INTO conversion_history 
        (username, filename, output_format, file_size, timestamp, processing_time, file_pages, batch_id, conversion_status, error_message)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (username, filename, output_format, file_size, datetime.now(), 
          processing_time, file_pages, batch_id, status, error))
    
    conn.commit()
    conn.close()

def get_user_history(username, limit=100):
    """Get user history with pagination"""
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("""
        SELECT * FROM conversion_history 
        WHERE username=? 
        ORDER BY timestamp DESC 
        LIMIT ?
    """, (username, limit))
    history = c.fetchall()
    conn.close()
    return history

def get_user_preferences(username):
    """Get user preferences"""
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("SELECT * FROM user_preferences WHERE username=?", (username,))
    prefs = c.fetchone()
    conn.close()
    
    if prefs:
        return {
            "default_output_format": prefs[1],
            "default_extraction_mode": prefs[2],
            "auto_convert": prefs[3],
            "notify_on_completion": prefs[4],
            "theme": prefs[5]
        }
    return {
        "default_output_format": "Excel (XLSX)",
        "default_extraction_mode": "Smart (Text + Tables)",
        "auto_convert": False,
        "notify_on_completion": True,
        "theme": "light"
    }

def save_user_preferences(username, preferences):
    """Save user preferences"""
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("""
        INSERT OR REPLACE INTO user_preferences 
        (username, default_output_format, default_extraction_mode, auto_convert, notify_on_completion, theme)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (username, 
          preferences.get("default_output_format", "Excel (XLSX)"),
          preferences.get("default_extraction_mode", "Smart (Text + Tables)"),
          preferences.get("auto_convert", False),
          preferences.get("notify_on_completion", True),
          preferences.get("theme", "light")))
    conn.commit()
    conn.close()

def log_user_activity(username, action, details=None):
    """Log user activity for audit trail"""
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    c.execute("""
        INSERT INTO user_activity (username, action, details, timestamp)
        VALUES (?, ?, ?, ?)
    """, (username, action, details, datetime.now()))
    conn.commit()
    conn.close()

# ============================================
# EXISTING FUNCTIONS (Keep all original processing functions)
# ============================================
# [All existing functions remain the same: 
#  extract_pdf_intelligent, convert_to_excel, convert_to_word, 
#  convert_to_markdown, convert_to_html, merge_pdfs, split_pdf, rotate_pdf]

def extract_pdf_intelligent(pdf_path, mode, extract_tables_flag):
    content = {"text": [], "tables": [], "pages": 0, "metadata": {}}
    
    try:
        doc = fitz.open(pdf_path)
        content["metadata"] = {
            "title": doc.metadata.get("title", "Unknown"),
            "author": doc.metadata.get("author", "Unknown"),
            "subject": doc.metadata.get("subject", "Unknown"),
            "creator": doc.metadata.get("creator", "Unknown"),
            "producer": doc.metadata.get("producer", "Unknown"),
            "page_count": len(doc)
        }
        doc.close()
    except:
        content["metadata"] = {"title": "Unknown", "author": "Unknown", "page_count": 0}
    
    with pdfplumber.open(pdf_path) as pdf:
        content["pages"] = len(pdf.pages)
        
        for page_num, page in enumerate(pdf.pages, 1):
            if mode in ["Smart (Text + Tables)", "Text Only"]:
                text = page.extract_text() or ""
                content["text"].append({
                    "page": page_num, 
                    "content": text,
                    "word_count": len(text.split()),
                    "char_count": len(text)
                })
            
            if extract_tables_flag and mode in ["Smart (Text + Tables)", "Tables Only"]:
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables, 1):
                    if table and len(table) > 1:
                        cleaned_table = []
                        for row in table:
                            cleaned_row = [str(cell).strip() if cell else "" for cell in row]
                            if any(cleaned_row):
                                cleaned_table.append(cleaned_row)
                        
                        if cleaned_table:
                            content["tables"].append({
                                "page": page_num,
                                "table_id": table_idx,
                                "rows": len(cleaned_table),
                                "columns": len(cleaned_table[0]) if cleaned_table else 0,
                                "table": cleaned_table
                            })
    return content

def convert_to_excel(content, output_buffer, metadata_include=True):
    wb = Workbook()
    ws_summary = wb.active
    ws_summary.title = "Summary"
    
    row = 1
    ws_summary.cell(row=row, column=1, value="PDF Intelligence Report")
    ws_summary.cell(row=row, column=1).font = Font(bold=True, size=14, color="D71E28")
    row += 2
    
    if metadata_include and content["metadata"]:
        ws_summary.cell(row=row, column=1, value="Document Information")
        ws_summary.cell(row=row, column=1).font = Font(bold=True, size=12)
        row += 1
        for key, value in content["metadata"].items():
            ws_summary.cell(row=row, column=1, value=key.replace('_', ' ').title())
            ws_summary.cell(row=row, column=2, value=str(value))
            row += 1
        row += 1
    
    ws_summary.cell(row=row, column=1, value="Statistics")
    ws_summary.cell(row=row, column=1).font = Font(bold=True, size=12)
    row += 1
    ws_summary.cell(row=row, column=1, value="Total Pages")
    ws_summary.cell(row=row, column=2, value=content['pages'])
    row += 1
    ws_summary.cell(row=row, column=1, value="Total Tables")
    ws_summary.cell(row=row, column=2, value=len(content['tables']))
    
    if content["text"]:
        ws_content = wb.create_sheet("Content")
        row = 1
        for page_data in content["text"]:
            ws_content.cell(row=row, column=1, value=f"--- PAGE {page_data['page']} ---")
            ws_content.cell(row=row, column=1).font = Font(bold=True, color="D71E28")
            row += 1
            for line in page_data["content"].split('\n'):
                if line.strip():
                    ws_content.cell(row=row, column=1, value=line)
                    row += 1
            row += 1
    
    if content["tables"]:
        for table_data in content["tables"]:
            sheet_name = f"Table_{table_data['table_id']}_P{table_data['page']}"[:31]
            ws_table = wb.create_sheet(sheet_name)
            table = table_data["table"]
            for r, row_data in enumerate(table, 1):
                for c, cell in enumerate(row_data, 1):
                    ws_table.cell(row=r, column=c, value=cell)
    
    wb.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer

def convert_to_word(content, output_buffer):
    doc = Document()
    doc.add_heading('PDF Intelligence Report', 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    if content["text"]:
        for page_data in content["text"]:
            doc.add_heading(f'Page {page_data["page"]}', level=1)
            doc.add_paragraph(page_data["content"])
    
    if content["tables"]:
        doc.add_heading('Extracted Tables', level=1)
        for table_data in content["tables"]:
            if table_data["table"]:
                doc.add_heading(f'Table {table_data["table_id"]} (Page {table_data["page"]})', level=2)
                table = doc.add_table(rows=len(table_data["table"]), cols=len(table_data["table"][0]))
                table.style = 'Light List Accent 1'
                for i, row in enumerate(table_data["table"]):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell
    
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer

def convert_to_markdown(content):
    md = []
    md.append("# PDF Intelligence Report\n\n")
    md.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
    
    for page_data in content["text"]:
        md.append(f"## Page {page_data['page']}\n\n")
        md.append(f"{page_data['content']}\n\n---\n\n")
    
    if content["tables"]:
        md.append("\n## Tables\n\n")
        for table_data in content["tables"]:
            md.append(f"### Table {table_data['table_id']}\n\n")
            if table_data["table"]:
                for i, row in enumerate(table_data["table"]):
                    md.append("| " + " | ".join(row) + " |\n")
                    if i == 0:
                        md.append("|" + "|".join(["---" for _ in row]) + "|\n")
            md.append("\n\n")
    
    return "".join(md)

def convert_to_html(content):
    html = []
    html.append(f"""<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"><title>PDF Intelligence Report</title>
<style>
body{{font-family:sans-serif;margin:40px}}
h1{{color:#d71e28}}
table{{border-collapse:collapse;width:100%}}
th,td{{border:1px solid #ddd;padding:8px}}
th{{background-color:#FDE8E8}}
</style></head>
<body>
<h1>PDF Intelligence Report</h1>
<p><strong>Generated:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>""")
    
    for page_data in content["text"]:
        html.append(f'<h2>Page {page_data["page"]}</h2>')
        html.append(f'<pre>{page_data["content"]}</pre>')
    
    if content["tables"]:
        html.append("<h2>Tables</h2>")
        for table_data in content["tables"]:
            html.append(f'<h3>Table {table_data["table_id"]} (Page {table_data["page"]})</h3>')
            html.append('<table>')
            if table_data["table"]:
                for i, row in enumerate(table_data["table"]):
                    html.append('<tr>')
                    tag = 'th' if i == 0 else 'td'
                    for cell in row:
                        html.append(f'<{tag}>{cell}</{tag}>')
                    html.append('</tr>')
            html.append('</table><br>')
    
    html.append("</body></html>")
    return "\n".join(html)

def merge_pdfs(pdf_files):
    merged = fitz.open()
    for pdf_file in pdf_files:
        pdf_content = pdf_file.read()
        with fitz.open(stream=pdf_content, filetype="pdf") as pdf:
            merged.insert_pdf(pdf)
    output = io.BytesIO()
    merged.save(output)
    merged.close()
    output.seek(0)
    return output

def split_pdf(pdf_bytes, pages_per_split):
    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    splits = []
    for i in range(0, len(pdf_doc), pages_per_split):
        new_pdf = fitz.open()
        end = min(i + pages_per_split, len(pdf_doc))
        new_pdf.insert_pdf(pdf_doc, from_page=i, to_page=end-1)
        output = io.BytesIO()
        new_pdf.save(output)
        new_pdf.close()
        output.seek(0)
        splits.append(output)
    pdf_doc.close()
    return splits

def rotate_pdf(pdf_bytes, rotation):
    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in pdf_doc:
        page.set_rotation(rotation)
    output = io.BytesIO()
    pdf_doc.save(output)
    pdf_doc.close()
    output.seek(0)
    return output

# ============================================
# STREAMLIT PAGE SETUP
# ============================================

# Initialize session state
def init_session_state():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if "session_id" not in st.session_state:
        st.session_state.session_id = hashlib.sha256(str(time.time()).encode()).hexdigest()[:16]
    if "last_activity" not in st.session_state:
        st.session_state.last_activity = datetime.now()
    if "username" not in st.session_state:
        st.session_state.username = None
    if "page" not in st.session_state:
        st.session_state.page = "📄 Convert PDF"
    if "theme" not in st.session_state:
        st.session_state.theme = "light"
    if "detected_lang" not in st.session_state:
        st.session_state.detected_lang = "eng"

def get_theme_colors():
    """Return theme colors based on current theme"""
    if st.session_state.theme == "dark":
        return {
            "bg": "#0f0f0f",
            "panel": "#1a1a1a",
            "panel2": "#1e1e1e",
            "text": "#ffffff",
            "muted": "#a0a0a0",
            "border": "rgba(255,255,255,0.10)",
            "border2": "rgba(255,255,255,0.14)",
            "accent": "#ff4444",
            "accent2": "#cc0000",
        }
    else:
        return {
            "bg": "#ffffff",
            "panel": "#ffffff",
            "panel2": "#f7f7f7",
            "text": "#111111",
            "muted": "#5b5b5b",
            "border": "rgba(0,0,0,0.10)",
            "border2": "rgba(0,0,0,0.14)",
            "accent": "#d71e28",
            "accent2": "#b5161f",
        }

def apply_style():
    t = get_theme_colors()
    is_dark = st.session_state.theme == "dark"
    
    st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&display=swap');

    * {{
        font-family: 'Inter', sans-serif !important;
    }}

    html, body, [data-testid="stAppViewContainer"], .stApp {{
        background: {t['bg']} !important;
        color: {t['text']} !important;
    }}

    #MainMenu, footer {{
        visibility: hidden;
    }}

    [data-testid="stHeader"] {{
        background: transparent !important;
    }}

    .block-container {{
        max-width: 1200px !important;
        padding-top: 2rem !important;
        padding-bottom: 2rem !important;
        padding-left: 2rem !important;
        padding-right: 2rem !important;
        margin: 0 auto !important;
    }}

    section[data-testid="stSidebar"] {{
        background: {t['panel']} !important;
        border-right: 1px solid {t['border']} !important;
    }}

    .notification-badge {{
        position: relative;
        display: inline-block;
    }}

    .notification-count {{
        position: absolute;
        top: -8px;
        right: -8px;
        background: {t['accent']};
        color: white;
        border-radius: 50%;
        padding: 2px 6px;
        font-size: 11px;
        font-weight: bold;
    }}

    .card {{
        background: {t['panel']};
        border: 1px solid {t['border']};
        border-radius: 18px;
        padding: 24px;
        margin-bottom: 20px;
    }}

    .card-soft {{
        background: {t['panel2']};
        border: 1px solid {t['border']};
        border-radius: 18px;
        padding: 24px;
        margin-bottom: 20px;
    }}

    .hero {{
        border: 1px solid {t['border']};
        border-radius: 22px;
        padding: 32px 24px;
        margin-bottom: 24px;
        background: {'radial-gradient(900px 260px at 50% -10%, rgba(255,68,68,0.15), transparent 60%)' if is_dark else 'radial-gradient(900px 260px at 50% -10%, rgba(215,30,40,0.10), transparent 60%)'};
    }}

    .title {{
        font-size: 28px;
        font-weight: 800;
        color: {t['text']};
    }}

    .login-title {{
        font-size: 24px;
        font-weight: 800;
        color: {t['text']};
    }}

    .subtitle {{
        margin-top: 8px;
        color: {t['muted']};
        font-size: 14px;
        line-height: 1.6;
    }}

    .chip {{
        display: inline-flex;
        align-items: center;
        gap: 8px;
        padding: 8px 14px;
        border-radius: 999px;
        border: 1px solid {t['border']};
        font-size: 12px;
        color: {t['muted']};
        background: {t['panel']};
        white-space: nowrap;
    }}

    .chip-dot {{
        width: 7px;
        height: 7px;
        border-radius: 50%;
        background: {t['accent']};
    }}

    .chip-container {{
        display: flex;
        gap: 10px;
        flex-wrap: wrap;
        justify-content: center;
        margin-top: 16px;
    }}

    .metric {{
        border: 1px solid {t['border']};
        border-radius: 18px;
        padding: 20px;
        background: {t['panel']};
    }}

    .metric-k {{
        font-size: 11px;
        color: {t['muted']};
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 1px;
        margin-bottom: 8px;
    }}

    .metric-v {{
        font-size: 28px;
        font-weight: 850;
        color: {t['text']};
    }}

    .muted {{
        color: {t['muted']};
    }}

    div.stButton > button {{
        background: {t['accent']} !important;
        color: white !important;
        border-radius: 12px !important;
        border: none !important;
        padding: 12px 20px !important;
        font-weight: 700 !important;
        font-size: 15px !important;
        min-height: 44px !important;
        width: 100% !important;
        transition: all 0.2s ease !important;
    }}

    div.stButton > button:hover {{
        background: {t['accent2']} !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15) !important;
    }}

    div[data-baseweb="input"],
    div[data-baseweb="select"] {{
        border-radius: 12px !important;
        border: 1px solid {t['border2']} !important;
        background: {t['panel']} !important;
    }}

    div[data-baseweb="input"] input,
    div[data-baseweb="select"] input {{
        color: {t['text']} !important;
        padding: 12px !important;
        font-size: 14px !important;
    }}

    [data-testid="stFileUploader"] {{
        border: 2px dashed {t['border2']} !important;
        border-radius: 16px !important;
        padding: 32px 24px !important;
        background: {t['panel']} !important;
        text-align: center !important;
        transition: all 0.2s ease !important;
        margin-bottom: 20px !important;
    }}

    [data-testid="stFileUploader"]:hover {{
        border-color: {t['accent']} !important;
        background: {'rgba(255,68,68,0.05)' if is_dark else 'rgba(215,30,40,0.02)'} !important;
    }}

    details {{
        border-radius: 14px !important;
        border: 1px solid {t['border']} !important;
        padding: 16px 20px !important;
        background: {t['panel']} !important;
        margin: 20px 0 !important;
    }}

    details summary {{
        font-weight: 700 !important;
        font-size: 15px !important;
        cursor: pointer !important;
        padding: 8px 0 !important;
        color: {t['text']} !important;
    }}

    .stTabs [data-baseweb="tab"] {{
        border-radius: 12px !important;
        padding: 12px 20px !important;
        font-weight: 700 !important;
        font-size: 14px !important;
        margin-right: 8px !important;
        background: {t['panel']} !important;
        border: 1px solid {t['border']} !important;
        color: {t['muted']} !important;
    }}

    .stTabs [aria-selected="true"] {{
        background: {'rgba(255,68,68,0.15)' if is_dark else 'rgba(215,30,40,0.10)'} !important;
        border: 1px solid {t['accent']} !important;
        color: {t['accent']} !important;
    }}

    @media (max-width: 768px) {{
        .title {{
            font-size: 22px;
        }}
        .login-title {{
            font-size: 20px;
        }}
        .block-container {{
            padding-left: 1rem !important;
            padding-right: 1rem !important;
        }}
        .hero {{
            padding: 24px 16px;
        }}
    }}
    </style>
    """, unsafe_allow_html=True)

def verify_user(username, password):
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    password_hash = hashlib.sha256(password.encode()).hexdigest()
    c.execute("SELECT * FROM users WHERE username=? AND password_hash=?", 
             (username, password_hash))
    user = c.fetchone()
    conn.close()
    return user is not None

def register_user(username, password, role='user'):
    conn = sqlite3.connect('users.db')
    c = conn.cursor()
    password_hash = hashlib.sha256(password.encode()).hexdigest()
    try:
        c.execute("INSERT INTO users (username, password_hash, role, created_at) VALUES (?, ?, ?, ?)",
                 (username, password_hash, role, datetime.now()))
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        return False
    finally:
        conn.close()

def sign_in(username, password):
    if verify_user(username, password):
        st.session_state.authenticated = True
        st.session_state.username = username
        st.session_state.last_activity = datetime.now()
        log_user_activity(username, "login")
        notification_system.add(f"Welcome back, {username}!", "success", duration=3)
        return True
    return False

def logout():
    if st.session_state.authenticated:
        log_user_activity(st.session_state.username, "logout")
    notification_system.add("You have been logged out", "info", duration=2)
    for k in list(st.session_state.keys()):
        del st.session_state[k]
    st.rerun()

def safe_rerun():
    try:
        st.rerun()
    except:
        pass

def toggle_theme():
    st.session_state.theme = "dark" if st.session_state.theme == "light" else "light"
    save_user_preferences(st.session_state.username, {"theme": st.session_state.theme})
    safe_rerun()

def touch():
    st.session_state.last_activity = datetime.now()

def is_timed_out():
    last = st.session_state.get("last_activity")
    if not last:
        return False
    return (datetime.now() - last).total_seconds() > SESSION_TIMEOUT_MINUTES * 60

# ============================================
# MAIN APP
# ============================================

st.set_page_config(
    page_title=f"{APP_NAME}",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded",
)

init_session_state()
apply_style()

# Initialize components
mobile_optimizer = MobileOptimizer()
batch_processor = BatchProcessor()
ocr_processor = OCRProcessor()

# Authentication
if not st.session_state.authenticated:
    theme_icon = "🌙" if st.session_state.theme == "light" else "☀️"
    st.markdown(f'<div style="position:fixed;top:20px;right:20px;z-index:9999;">', unsafe_allow_html=True)
    if st.button(f"{theme_icon} Theme", key="theme_toggle_login"):
        toggle_theme()
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown('<div style="height: 1.8rem;"></div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1, 1.25, 1])
    with c2:
        st.markdown(f"""
        <div class="card" style="margin-top: 10vh;">
            <div class="login-title" style="text-align:center;">⚡ {APP_NAME}</div>
            <div class="subtitle" style="text-align:center;">Sign in to continue.</div>
            <div style="height: 14px;"></div>
            <div style="display:flex; justify-content:center;">
                <div class="chip"><span class="chip-dot"></span> Version {APP_VERSION} • Production</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        tab1, tab2 = st.tabs(["Sign In", "Register"])
        
        with tab1:
            with st.form("login_form", clear_on_submit=True):
                username = st.text_input("Username", placeholder="Enter username")
                password = st.text_input("Password", type="password", placeholder="Enter password")
                col1, col2 = st.columns(2)
                with col1:
                    ok = st.form_submit_button("Sign in", use_container_width=True)
                with col2:
                    demo = st.form_submit_button("Demo", use_container_width=True)

            if ok:
                if username and password:
                    if sign_in(username, password):
                        st.success("Sign in successful!")
                        safe_rerun()
                    else:
                        st.error("Invalid credentials")
                else:
                    st.warning("Please fill in all fields")
            
            if demo:
                st.info("Demo: admin / admin123")
        
        with tab2:
            with st.form("register_form", clear_on_submit=True):
                new_username = st.text_input("New Username", placeholder="Choose username")
                new_password = st.text_input("New Password", type="password", placeholder="Min 6 characters")
                confirm_password = st.text_input("Confirm Password", type="password", placeholder="Confirm password")
                reg = st.form_submit_button("Register", use_container_width=True)

            if reg:
                if new_username and new_password and confirm_password:
                    if new_password == confirm_password:
                        if len(new_password) >= 6:
                            if register_user(new_username, new_password):
                                st.success("Registration successful! Please sign in.")
                            else:
                                st.error("Username already exists")
                        else:
                            st.warning("Password must be at least 6 characters")
                    else:
                        st.error("Passwords don't match")
                else:
                    st.warning("Please fill in all fields")
    st.stop()

# Session timeout check
if st.session_state.authenticated and is_timed_out():
    st.session_state.authenticated = False
    st.warning("Session timed out. Sign in again.")
    safe_rerun()

touch()

# ============================================
# MAIN DASHBOARD
# ============================================

with st.sidebar:
    # Header with notification badge
    unread_count = notification_system.get_unread_count()
    badge_html = f'<span class="notification-count">{unread_count}</span>' if unread_count > 0 else ''
    
    st.markdown(f"""
    <div style="display: flex; justify-content: space-between; align-items: center;">
        <div>
            <strong>⚡ {st.session_state.username}</strong>
        </div>
        <div class="notification-badge">
            🔔 {badge_html}
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # Theme toggle
    theme_icon = "🌙 Dark Mode" if st.session_state.theme == "light" else "☀️ Light Mode"
    if st.button(theme_icon, use_container_width=True, key="theme_toggle_sidebar"):
        toggle_theme()
    
    st.markdown("---")
    
    # Navigation
    page = st.radio("Navigation", [
        "📄 Convert PDF", 
        "📦 Batch Convert",
        "🔍 OCR Scanner",
        "🔧 PDF Tools", 
        "📊 History", 
        "🔔 Notifications",
        "⚙️ Settings"
    ], label_visibility="collapsed")
    st.session_state.page = page
    
    st.markdown("---")
    
    # Recent activity
    history = get_user_history(st.session_state.username, limit=5)
    if history:
        st.markdown("**Recent Activity**")
        for h in history[:5]:
            st.markdown(f"📄 {str(h[2])[:30]} → {h[3]}")
    
    st.markdown("---")
    
    if st.button("🚪 Sign Out", use_container_width=True):
        logout()
    
    # Mobile notice
    mobile_optimizer.render_mobile_notice()

# Main content area
st.markdown(f"""
<div class="hero" style="text-align:center;">
    <div class="title">⚡ {APP_NAME}</div>
    <div class="subtitle">Upload your PDF document. Extract content, convert formats, and manage your documents intelligently.</div>
    <div class="chip-container">
        <div class="chip"><span class="chip-dot"></span> Secure session</div>
        <div class="chip">Session {st.session_state.session_id}</div>
        <div class="chip">Production</div>
        <div class="chip">Version {APP_VERSION}</div>
        <div class="chip">User {st.session_state.username}</div>
    </div>
</div>
""", unsafe_allow_html=True)

# Page routing
if page == "📄 Convert PDF":
    # Single PDF conversion (original implementation)
    st.markdown("""<div class="card"><div style="font-size:16px; font-weight:800;">Document Conversion</div><div class="subtitle">Upload your PDF for intelligent extraction and multi-format conversion.</div></div>""", unsafe_allow_html=True)
    st.markdown('<div class="section-spacer"></div>', unsafe_allow_html=True)
    
    col_input, col_settings = st.columns([2, 1])
    
    with col_input:
        uploaded_file = st.file_uploader("Choose a PDF file", type=['pdf'], key="convert_upload")
        if uploaded_file:
            st.markdown(f"""<div class="card-soft"><strong>📄 Document:</strong> {uploaded_file.name}<br><strong>📏 Size:</strong> {len(uploaded_file.getvalue())/1024:.2f} KB</div>""", unsafe_allow_html=True)
    
    with col_settings:
        output_format = st.selectbox("Convert to", ["Excel (XLSX)", "Word (DOCX)", "Text (TXT)", "CSV", "JSON", "Markdown", "HTML"])
        extraction_modes = ["Smart (Text + Tables)", "Text Only", "Tables Only"]
        if OCR_AVAILABLE:
            extraction_modes.append("OCR (Scanned PDFs)")
        extraction_mode = st.selectbox("Extraction mode", extraction_modes)
        extract_tables = st.checkbox("Extract tables", value=True)
    
    st.markdown('<div class="section-spacer"></div>', unsafe_allow_html=True)
    
    with st.expander("🔧 Advanced Options"):
        include_metadata = st.checkbox("Include metadata", value=True)
    
    st.markdown('<div class="section-spacer"></div>', unsafe_allow_html=True)
    
    b1, b2 = st.columns([1, 5])
    with b1:
        convert_button = st.button("🔄 Convert", use_container_width=True)
    with b2:
        if st.button("Clear", use_container_width=True):
            st.rerun()
    
    if uploaded_file and convert_button:
        with st.spinner("Processing..."):
            try:
                start_time = time.time()
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_path = tmp_file.name
                
                content = extract_pdf_intelligent(tmp_path, extraction_mode, extract_tables)
                processing_time = time.time() - start_time
                
                result_buffer = None
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                base_name = Path(uploaded_file.name).stem
                
                if output_format == "Excel (XLSX)":
                    result_buffer = io.BytesIO()
                    convert_to_excel(content, result_buffer, include_metadata)
                    filename = f"{base_name}_{timestamp}.xlsx"
                    mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                elif output_format == "Word (DOCX)":
                    result_buffer = io.BytesIO()
                    convert_to_word(content, result_buffer)
                    filename = f"{base_name}_{timestamp}.docx"
                    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                elif output_format == "Text (TXT)":
                    text_content = "\n\n".join([p["content"] for p in content["text"]])
                    result_buffer = io.BytesIO(text_content.encode('utf-8'))
                    filename = f"{base_name}_{timestamp}.txt"
                    mime = "text/plain"
                elif output_format == "CSV":
                    result_buffer = io.BytesIO()
                    text_stream = io.TextIOWrapper(result_buffer, 'utf-8', newline='')
                    writer = csv.writer(text_stream)
                    writer.writerow(["Page", "Content"])
                    for page_data in content["text"]:
                        writer.writerow([page_data["page"], page_data["content"]])
                    text_stream.flush()
                    result_buffer.seek(0)
                    filename = f"{base_name}_{timestamp}.csv"
                    mime = "text/csv"
                elif output_format == "JSON":
                    json_content = {"metadata": content["metadata"], "pages": content["pages"], "text": [{"page": p["page"], "content": p["content"]} for p in content["text"]], "tables": [{"page": t["page"], "table": t["table"]} for t in content["tables"]]}
                    result_buffer = io.BytesIO(json.dumps(json_content, indent=2, ensure_ascii=False).encode('utf-8'))
                    filename = f"{base_name}_{timestamp}.json"
                    mime = "application/json"
                elif output_format == "Markdown":
                    md = convert_to_markdown(content)
                    result_buffer = io.BytesIO(md.encode('utf-8'))
                    filename = f"{base_name}_{timestamp}.md"
                    mime = "text/markdown"
                elif output_format == "HTML":
                    html_content = convert_to_html(content)
                    result_buffer = io.BytesIO(html_content.encode('utf-8'))
                    filename = f"{base_name}_{timestamp}.html"
                    mime = "text/html"
                
                os.unlink(tmp_path)
                
                # Save to history with enhanced data
                save_conversion_history(
                    st.session_state.username, 
                    uploaded_file.name, 
                    output_format, 
                    len(uploaded_file.getvalue()),
                    processing_time,
                    content["pages"]
                )
                
                notification_system.add(
                    f"Successfully converted {uploaded_file.name} to {output_format}",
                    "success",
                    duration=3
                )
                
                st.markdown("")
                m1, m2, m3 = st.columns(3)
                with m1:
                    st.markdown(f"<div class='metric'><div class='metric-k'>Pages</div><div class='metric-v'>{content['pages']}</div></div>", unsafe_allow_html=True)
                with m2:
                    st.markdown(f"<div class='metric'><div class='metric-k'>Tables</div><div class='metric-v'>{len(content['tables'])}</div></div>", unsafe_allow_html=True)
                with m3:
                    total_words = sum(t.get("word_count", 0) for t in content["text"])
                    st.markdown(f"<div class='metric'><div class='metric-k'>Words</div><div class='metric-v'>{total_words}</div></div>", unsafe_allow_html=True)
                
                st.success(f"Conversion complete! ({output_format})")
                
                if result_buffer:
                    result_buffer.seek(0)
                    st.download_button(label=f"💾 Download {filename}", data=result_buffer.getvalue(), file_name=filename, mime=mime, use_container_width=True)
                
                st.balloons()
                
            except Exception as e:
                st.error(f"Error: {str(e)}")
                notification_system.add(f"Conversion failed: {str(e)}", "error", duration=3)

elif page == "📦 Batch Convert":
    st.markdown("""<div class="card"><div style="font-size:16px; font-weight:800;">Batch Processing</div><div class="subtitle">Convert multiple PDF files at once. Save time and effort.</div></div>""", unsafe_allow_html=True)
    st.markdown("")
    
    # Batch upload
    uploaded_files = st.file_uploader(
        f"Upload PDF files (Max {MAX_BATCH_SIZE} files)", 
        type=['pdf'], 
        accept_multiple_files=True,
        key="batch_upload"
    )
    
    if uploaded_files:
        if len(uploaded_files) > MAX_BATCH_SIZE:
            st.warning(f"Maximum {MAX_BATCH_SIZE} files allowed. Please reduce selection.")
        else:
            st.info(f"📁 {len(uploaded_files)} files selected. Total size: {sum(len(f.getvalue()) for f in uploaded_files)/1024/1024:.2f} MB")
            
            col1, col2 = st.columns(2)
            
            with col1:
                output_format = st.selectbox(
                    "Output Format",
                    ["Excel (XLSX)", "Word (DOCX)", "Text (TXT)", "CSV", "JSON", "Markdown", "HTML"],
                    key="batch_format"
                )
            
            with col2:
                extraction_mode = st.selectbox(
                    "Extraction Mode",
                    ["Smart (Text + Tables)", "Text Only", "Tables Only"],
                    key="batch_mode"
                )
            
            include_metadata = st.checkbox("Include metadata", value=True)
            extract_tables = st.checkbox("Extract tables", value=True)
            
            if st.button("🚀 Start Batch Conversion", use_container_width=True):
                batch_id = batch_processor.create_batch_job(
                    uploaded_files, 
                    output_format, 
                    extraction_mode,
                    {"include_metadata": include_metadata, "extract_tables": extract_tables}
                )
                
                # Process with progress
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                def update_progress(progress):
                    progress_bar.progress(int(progress))
                    status_text.text(f"Processing: {progress:.1f}% complete")
                
                batch_processor.process_batch(batch_id, update_progress)
                
                status_text.empty()
                progress_bar.empty()
                
                # Show results
                batch_processor.render_batch_status(batch_id)
    
    # Show active batch if exists
    if st.session_state.active_batch and st.session_state.active_batch in st.session_state.batch_jobs:
        st.markdown("---")
        st.markdown("### Current Batch Status")
        batch_processor.render_batch_status(st.session_state.active_batch)

elif page == "🔍 OCR Scanner":
    if OCR_AVAILABLE:
        ocr_processor.render_ocr_interface()
    else:
        st.error("OCR not available. Please install pytesseract and pdf2image.")
        st.markdown("""
        **To enable OCR, install:**
        ```bash
        pip install pytesseract pdf2image langdetect
